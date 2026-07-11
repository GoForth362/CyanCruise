from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("docs/knowledge-base/cyancruise升学规划知识库")
PACKAGE_NAME = "CyanCruise_升学规划知识库_资料包.zip"


SOURCES = {
    "研招网报名公告": "https://yz.chsi.com.cn/kyzx/other/202509/20250930/2293435905.html",
    "研招网首页": "https://yz.chsi.com.cn/",
    "研招网硕士目录": "https://yz.chsi.com.cn/zsml/",
    "研招网院校库": "https://yz.chsi.com.cn/sch/",
    "推免服务系统": "https://yz.chsi.com.cn/tm/",
    "教育部推免管理办法": "https://www.moe.gov.cn/srcsite/A15/moe_778/s3113/200607/t20060712_79975.html",
    "教育部完善推免工作通知": "https://www.moe.gov.cn/srcsite/A15/moe_778/s3261/201408/t20140804_172730.html",
    "中国教育考试网研考": "https://yankao.neea.edu.cn/",
    "教育部本科专业目录": "https://www.moe.gov.cn/srcsite/A08/moe_1034/s4930/202504/W020250422312780837078.pdf",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def style_table(table, widths=None):
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            if widths:
                set_cell_width(cell, widths[col_idx])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)
                for run in paragraph.runs:
                    run.font.name = "微软雅黑"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                    run.font.size = Pt(9.5)
            if row_idx == 0:
                set_cell_shading(cell, "F2F4F7")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def set_doc_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(20)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("1F4D78")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(subtitle)
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("555555")


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.1)
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.1)
        p.add_run(item)


def add_source_section(doc, source_names):
    doc.add_heading("资料来源", level=1)
    doc.add_paragraph("以下来源均为公开信息入口。具体年份的报名、考试、复试、推免安排，应以当年官方公告和招生单位通知为准。")
    for name in source_names:
        doc.add_paragraph(f"{name}：{SOURCES[name]}", style="List Bullet")


def save_doc(doc, filename):
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUT_DIR / filename
    doc.save(path)
    return path


def build_kaoyan_doc():
    doc = Document()
    set_doc_styles(doc)
    add_title(doc, "大学生考研规划指南", "CyanCruise 升学规划知识库资料")

    doc.add_heading("适用场景", level=1)
    doc.add_paragraph(
        "本文用于支持 CyanCruise 智能体回答“是否考研”“如何准备考研”“如何选择目标院校和专业”等问题。"
        "内容依据研招网、中国教育考试网等官方公开入口整理，适合放入升学规划知识库作为检索资料。"
    )

    doc.add_heading("官方流程要点", level=1)
    add_bullets(doc, [
        "硕士研究生招生考试通常包括网上报名、网上确认、初试、复试、调剂和录取等环节。",
        "网上报名和网上确认均应以研招网、省级教育招生考试机构、报考点和招生单位公告为准。",
        "考生通常只能保留一条有效报名信息，报名结束后应重点关注网上确认、准考证下载、初试安排和复试要求。",
        "考试时间、报名时间、确认方式、复试线、调剂规则每年可能调整，智能体回答时必须提示用户核对当年公告。",
    ])

    doc.add_heading("本科阶段准备节奏", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "阶段"
    table.rows[0].cells[1].text = "重点任务"
    table.rows[0].cells[2].text = "CyanCruise 建议"
    rows = [
        ("大一至大二", "夯实专业课、英语、数学或专业基础，保留课程成绩和项目记录。", "暂不急于定校，先判断自己是否适合学术深造或专业提升。"),
        ("大三上", "明确是否考研，初步确定专业方向、城市偏好和目标层次。", "建议建立“冲刺、稳妥、保底”三个目标层次。"),
        ("大三下至暑期", "进入系统复习，完成公共课、专业课第一轮和强化复习。", "每月复盘分数、进度、错题和心理压力，必要时调整院校层次。"),
        ("大四上", "完成报名、网上确认、准考证、模拟考试和冲刺复习。", "避免只看热度选校，要结合招生人数、考试科目、复试比例和个人基础。"),
        ("初试后", "准备复试、联系导师、整理材料，关注调剂和复试线。", "初试估分后同步准备复试与调剂备选，不等分数公布才行动。"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = value
    style_table(table, [1500, 4200, 3660])

    doc.add_heading("考研适配度判断", level=1)
    add_bullets(doc, [
        "目标动机：是为了学术研究、学历提升、专业转换，还是暂时逃避就业压力。",
        "基础能力：专业课、英语、数学或目标专业核心课是否能支撑一年左右高强度备考。",
        "时间条件：是否有稳定备考时间，是否需要同时承担实习、项目、竞赛或家庭压力。",
        "目标合理性：目标院校、专业、城市是否与个人成绩、复习基础、风险承受能力匹配。",
        "备选方案：是否准备就业、调剂、二战、留学或其他路径，避免单一路径失败造成被动。",
    ])

    doc.add_heading("备考规划模板", level=1)
    add_numbers(doc, [
        "建立用户画像：专业、年级、GPA、排名、英语/数学基础、目标城市、专业兴趣、家庭约束。",
        "确定专业方向：优先判断本专业深造、跨专业、专业学位或学术学位的匹配度。",
        "筛选目标院校：从研招网硕士目录、院校库、招生简章、往年复试线和招生人数交叉验证。",
        "拆解科目计划：公共课按基础、强化、真题、模拟四个阶段推进；专业课按教材、真题、专题、背诵或案例训练推进。",
        "设置复盘节点：每两到四周检查进度、正确率、薄弱科目和情绪状态，必要时调整目标。",
    ])

    doc.add_heading("智能体回答边界", level=1)
    add_bullets(doc, [
        "可以给出规划建议、风险提醒和资料核对路径。",
        "不得编造某院校当年招生人数、分数线、复试比例或考试科目。",
        "对具体年份安排应说明“以当年研招网、省级考试机构和招生单位公告为准”。",
        "当用户信息不足时，应优先追问专业、年级、GPA/排名、目标方向、英语/数学基础和可用时间。",
    ])

    add_source_section(doc, ["研招网报名公告", "研招网首页", "研招网硕士目录", "研招网院校库", "中国教育考试网研考"])
    return save_doc(doc, "大学生考研规划指南.docx")


def build_tuimian_doc():
    doc = Document()
    set_doc_styles(doc)
    add_title(doc, "保研条件与准备材料", "CyanCruise 升学规划知识库资料")

    doc.add_heading("适用场景", level=1)
    doc.add_paragraph(
        "本文用于支持 CyanCruise 智能体回答“我能不能保研”“保研要准备什么材料”“保研和考研如何选择”等问题。"
        "内容依据教育部推免政策、研招网推免服务系统等公开资料整理。"
    )

    doc.add_heading("基本概念", level=1)
    doc.add_paragraph(
        "推荐免试攻读研究生通常面向获得所在高校推免资格的优秀应届本科毕业生。"
        "推免不参加全国硕士研究生招生考试初试，但仍需参加招生单位组织的复试或考核。"
    )

    doc.add_heading("常见资格维度", level=1)
    add_bullets(doc, [
        "学业成绩：课程成绩、GPA、专业排名通常是基础条件，具体规则由所在高校确定。",
        "综合表现：科研训练、竞赛获奖、社会实践、学生工作、志愿服务等可能纳入综合评价。",
        "诚信与纪律：学术诚信、考试纪律、在校处分等会影响资格认定。",
        "外语和科研潜力：部分招生单位会关注英语能力、科研经历、论文、专利、项目经历或面试表现。",
        "专项计划或特殊类别：以高校、招生单位和教育主管部门当年政策为准。",
    ])

    doc.add_heading("材料清单", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "材料类别"
    table.rows[0].cells[1].text = "常见内容"
    table.rows[0].cells[2].text = "准备建议"
    rows = [
        ("身份与学籍", "身份证、学生证、学籍证明、推免资格证明。", "提前核对学信网学籍信息，确保姓名、证件号、学校信息一致。"),
        ("成绩材料", "成绩单、GPA/排名证明、专业排名说明。", "使用学校教务部门盖章版本，避免自行截图替代正式材料。"),
        ("能力证明", "英语四六级、雅思托福、竞赛、论文、专利、科研项目、实习实践。", "按“名称、时间、角色、成果、证明材料”整理。"),
        ("申请文书", "个人陈述、简历、研究计划、导师联系邮件。", "突出专业兴趣、研究经历和未来计划，不夸大经历。"),
        ("复试材料", "自我介绍、专业问题准备、项目问答、综合面试材料。", "围绕目标专业准备三到五个核心问题。"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = value
    style_table(table, [1700, 4100, 3560])

    doc.add_heading("时间线建议", level=1)
    add_numbers(doc, [
        "大一至大二：稳定成绩排名，积累科研、竞赛、项目和实践记录。",
        "大三上：确认本校推免规则，评估排名、课程短板和加分项。",
        "大三下至暑期：准备简历、个人陈述、成绩单、证明材料，关注目标院校夏令营和预推免。",
        "九月至十月：关注推免服务系统和招生单位通知，按要求填报志愿、接收复试和待录取通知。",
        "拟录取后：核对系统确认状态、学校公示和后续学籍注册要求。",
    ])

    doc.add_heading("智能体判断逻辑", level=1)
    add_bullets(doc, [
        "如果用户排名靠前且有科研/竞赛/项目积累，可优先建议评估保研资格。",
        "如果用户成绩边缘但综合经历较好，应建议查阅本校推免细则，不能直接判断一定可保研。",
        "如果用户已错过本校推免资格认定，应转向考研、就业、留学或其他升学路径比较。",
        "如果用户只提供 GPA 而没有排名，应追问专业排名、学院规则和是否有加分项。",
    ])

    doc.add_heading("注意事项", level=1)
    add_bullets(doc, [
        "推免资格由推荐高校按规定产生，接收结果由招生单位考核确定。",
        "推免相关系统、时间、志愿规则每年可能调整，应以当年推免服务系统和招生单位通知为准。",
        "材料必须真实准确。对科研、竞赛、项目经历的描述应能被证明材料支持。",
    ])

    add_source_section(doc, ["推免服务系统", "教育部推免管理办法", "教育部完善推免工作通知", "研招网首页"])
    return save_doc(doc, "保研条件与准备材料.docx")


def build_school_major_doc():
    doc = Document()
    set_doc_styles(doc)
    add_title(doc, "院校专业选择方法", "CyanCruise 升学规划知识库资料")

    doc.add_heading("适用场景", level=1)
    doc.add_paragraph(
        "本文用于支持 CyanCruise 智能体帮助学生选择考研、保研或其他升学目标时，"
        "从院校层次、专业方向、城市、考试科目、竞争强度和个人画像进行综合匹配。"
    )

    doc.add_heading("信息核验入口", level=1)
    add_bullets(doc, [
        "研招网硕士目录：用于查询当年硕士招生专业、研究方向、学习方式、考试科目等。",
        "研招网院校库：用于查询招生单位、所在地、主管部门和院校特性等基本信息。",
        "招生单位研究生院官网：用于核对招生简章、专业目录、复试办法、导师与学院通知。",
        "教育部本科专业目录：用于识别本科专业门类和专业归属，辅助跨专业判断。",
    ])

    doc.add_heading("选择院校和专业的六个维度", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "维度"
    table.rows[0].cells[1].text = "看什么"
    table.rows[0].cells[2].text = "智能体应如何使用"
    rows = [
        ("专业匹配", "本科专业、目标专业、跨专业跨度、核心课程基础。", "先判断是本专业深造、相近专业转换还是大跨度跨考。"),
        ("院校层次", "学校平台、学科实力、招生单位特色、培养方向。", "不要只按名气推荐，要结合用户成绩和风险承受能力。"),
        ("城市与发展", "目标城市、产业机会、家庭约束、未来就业区域。", "升学目标应与长期发展城市和行业机会联动。"),
        ("考试科目", "公共课、专业课、统考或自命题、是否考数学。", "如果用户基础薄弱，应重点提示科目风险。"),
        ("竞争强度", "招生人数、复试线、报录趋势、复试比例、推免比例。", "必须提醒用户核对官方数据，不能凭经验给绝对结论。"),
        ("备选路径", "冲刺、稳妥、保底，调剂和就业备选。", "至少给出两到三个目标层次，避免单点押注。"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = value
    style_table(table, [1500, 3600, 4260])

    doc.add_heading("目标分层方法", level=1)
    add_bullets(doc, [
        "冲刺目标：平台或专业实力明显提升，但需要用户在复习、材料或面试上付出更高强度。",
        "稳妥目标：与用户当前成绩、基础和时间较匹配，是主攻层次。",
        "保底目标：风险较低，用于降低升学路径的不确定性。",
        "不建议目标：考试科目明显不匹配、跨专业跨度过大、招生人数极少或用户准备时间不足。",
    ])

    doc.add_heading("专业选择方法", level=1)
    add_numbers(doc, [
        "先确认用户本科专业门类和核心课程，再判断可延伸方向。",
        "确认用户动机：学术研究、就业提升、转行、城市发展或学历提升。",
        "用硕士目录查询目标专业的招生单位、考试科目和学习方式。",
        "查看招生单位官网，核对是否有特殊报考条件、复试要求或参考书。",
        "把目标专业与用户能力差距拆成课程、技能、材料、面试和时间五类任务。",
    ])

    doc.add_heading("智能体追问模板", level=1)
    add_bullets(doc, [
        "你的本科专业、年级、GPA 和专业排名大致是多少？",
        "你更看重学校平台、专业方向、城市，还是未来就业机会？",
        "你的英语、数学或目标专业核心课基础如何？",
        "你能接受跨专业备考吗？是否有相关课程、项目或实践经历？",
        "你希望选择冲刺型目标，还是更希望稳妥上岸？",
    ])

    doc.add_heading("知识库回答边界", level=1)
    add_bullets(doc, [
        "可提供院校专业选择框架、信息核验路径和目标分层建议。",
        "不得替代招生单位官方公告，不得虚构当年招生名额、复试线或考试科目。",
        "涉及具体年份和具体院校时，应提示用户以研招网和招生单位官网为准。",
    ])

    add_source_section(doc, ["研招网硕士目录", "研招网院校库", "研招网首页", "教育部本科专业目录"])
    return save_doc(doc, "院校专业选择方法.docx")


def main():
    paths = [build_kaoyan_doc(), build_tuimian_doc(), build_school_major_doc()]
    zip_path = OUT_DIR / PACKAGE_NAME
    with ZipFile(zip_path, "w", ZIP_DEFLATED) as zf:
        for path in paths:
            zf.write(path, path.name)
    print("Generated:")
    for path in paths:
        print(path.resolve())
    print(zip_path.resolve())


if __name__ == "__main__":
    main()
