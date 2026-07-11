from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("docs/knowledge-base/cyancruise简历面试知识库")
PACKAGE_NAME = "CyanCruise_简历面试知识库_资料包.zip"


SOURCES = {
    "国家大学生就业服务平台": "https://www.ncss.cn/",
    "国家大学生就业服务平台职位页": "https://www.ncss.cn/student/m/jobs/index.html",
    "教育部推广使用国家24365平台通知": "https://www.moe.gov.cn/srcsite/A15/s3265/202204/t20220406_614117.html",
    "中国公共招聘网": "https://job.mohrss.gov.cn/",
    "全国就业公共服务平台": "https://www.12333.gov.cn/job/?channel=12333",
    "就业在线": "https://www.jobonline.cn/",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def style_table(table, widths=None):
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            if widths:
                set_cell_width(cell, widths[col_idx])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)
                for run in paragraph.runs:
                    run.font.name = "微软雅黑"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                    run.font.size = Pt(9.5)
            if row_idx == 0:
                set_cell_shading(cell, "F2F4F7")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def set_doc_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(20)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("1F4D78")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(subtitle)
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("555555")


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.1)
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.1)
        p.add_run(item)


def add_source_section(doc, source_names):
    doc.add_heading("资料来源", level=1)
    doc.add_paragraph("以下来源为公开就业服务和招聘信息入口。简历与面试建议为 CyanCruise 面向大学生求职场景整理的通用方法，实际要求以岗位公告和用人单位通知为准。")
    for name in source_names:
        doc.add_paragraph(f"{name}：{SOURCES[name]}", style="List Bullet")


def save_doc(doc, filename):
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUT_DIR / filename
    doc.save(path)
    return path


def build_resume_guide():
    doc = Document()
    set_doc_styles(doc)
    add_title(doc, "大学生简历写作指南", "CyanCruise 简历面试知识库资料")

    doc.add_heading("适用场景", level=1)
    doc.add_paragraph(
        "本文用于支持 CyanCruise 智能体回答“简历怎么写”“没有实习怎么写简历”“如何让简历匹配岗位要求”等问题。"
        "内容面向大学生求职、实习、校招、科研助理和基层项目等场景。"
    )

    doc.add_heading("简历的核心目标", level=1)
    add_bullets(doc, [
        "让招聘方快速判断候选人是否符合岗位要求。",
        "用课程、项目、实习、竞赛、作品、证书等证据证明能力。",
        "突出与目标岗位最相关的信息，而不是罗列全部经历。",
        "保持真实、清楚、可核验，不夸大职责和成果。",
    ])

    doc.add_heading("推荐结构", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "模块"
    table.rows[0].cells[1].text = "写什么"
    table.rows[0].cells[2].text = "注意事项"
    rows = [
        ("基本信息", "姓名、电话、邮箱、城市、求职方向。", "邮箱和电话要准确，求职方向应与投递岗位一致。"),
        ("教育背景", "学校、专业、学历、时间、GPA/排名、核心课程。", "成绩有优势就写，课程只放与岗位相关内容。"),
        ("项目/实践", "项目背景、个人角色、关键任务、方法工具、结果。", "优先写成果和个人贡献，避免只写团队做了什么。"),
        ("实习经历", "公司/组织、岗位、时间、职责、产出、数据。", "用动词开头，突出解决的问题和实际影响。"),
        ("校园经历", "社团、学生工作、志愿服务、竞赛活动。", "只有与岗位能力相关时重点展开。"),
        ("技能证书", "工具、语言、证书、作品链接、资格考试。", "不要写“熟练”但没有证据支撑的技能。"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = value
    style_table(table, [1600, 3900, 3860])

    doc.add_heading("写作原则", level=1)
    add_numbers(doc, [
        "先读岗位要求，再决定简历重点。不同岗位应准备不同版本。",
        "经历表达遵循“场景、任务、行动、结果”的顺序，但对用户可见文案写作中称为“经历讲述框架”。",
        "尽量使用可量化结果，如人数、周期、准确率、转化率、成本、效率、作品数量等。",
        "把课程作业、毕业设计、竞赛、社团项目转化为可验证的能力证据。",
        "一页优先，信息密度适中，格式统一，避免花哨模板影响阅读。",
    ])

    doc.add_heading("常见问题与修改方向", level=1)
    add_bullets(doc, [
        "问题：只写“负责、参与、协助”。修改：写清具体任务、方法和结果。",
        "问题：项目描述太像课程介绍。修改：写个人角色、技术/方法选择和最终产出。",
        "问题：经历与岗位不匹配。修改：把目标岗位高频要求映射到简历证据。",
        "问题：没有实习。修改：用课程项目、竞赛、实验、志愿服务、作品集补充能力证明。",
        "问题：技能堆砌。修改：每项技能尽量对应一个项目、证书或作品。",
    ])

    doc.add_heading("智能体回答边界", level=1)
    add_bullets(doc, [
        "可以给出简历结构、修改建议和可替换表达。",
        "不能伪造实习、项目、奖项、证书或量化成果。",
        "当用户经历不足时，应建议补充真实项目、作品或实践，而不是夸大内容。",
        "用户可见文案避免使用 JD、STAR 等缩写，改为“岗位要求”“经历讲述框架”。",
    ])

    add_source_section(doc, ["国家大学生就业服务平台", "国家大学生就业服务平台职位页", "教育部推广使用国家24365平台通知"])
    return save_doc(doc, "大学生简历写作指南.docx")


def build_experience_examples():
    doc = Document()
    set_doc_styles(doc)
    add_title(doc, "实践经历表达示例", "CyanCruise 简历面试知识库资料")

    doc.add_heading("适用场景", level=1)
    doc.add_paragraph(
        "本文用于支持 CyanCruise 智能体把学生的项目、实习、竞赛、科研、社团和志愿经历改写成更适合简历和面试的表达。"
        "所有示例均为表达模板，不能替代用户真实经历。"
    )

    doc.add_heading("经历讲述框架", level=1)
    add_bullets(doc, [
        "场景：这段经历发生在什么项目、团队、课程、活动或岗位中。",
        "任务：你具体负责什么问题或目标。",
        "行动：你用了什么方法、工具、流程或沟通方式。",
        "结果：产生了什么产出、数据、反馈或沉淀。",
    ])

    doc.add_heading("表达示例", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "经历类型"
    table.rows[0].cells[1].text = "普通写法"
    table.rows[0].cells[2].text = "优化写法"
    rows = [
        ("课程项目", "参与课程项目，完成系统设计。", "在课程项目中负责需求梳理与核心模块实现，完成流程设计、功能测试和说明文档，支撑团队按期完成演示。"),
        ("科研训练", "参与老师课题，查阅资料。", "围绕课题方向检索并整理文献，归纳研究问题、方法和数据来源，形成阶段性综述材料供小组讨论。"),
        ("竞赛经历", "参加创新创业比赛并获奖。", "负责方案设计与路演材料制作，完成用户调研、竞品分析和商业模式梳理，帮助团队进入校级决赛。"),
        ("社团活动", "组织社团活动。", "统筹活动报名、物料、现场流程和复盘反馈，协调多名成员完成执行，提升活动参与度和组织效率。"),
        ("志愿服务", "参加志愿服务。", "面向服务对象完成接待、指引、信息登记和问题反馈，提升沟通耐心和现场应变能力。"),
        ("实习经历", "负责日常运营工作。", "根据运营目标整理用户反馈、维护内容排期并跟踪数据表现，形成周报，辅助团队调整活动节奏。"),
        ("设计作品", "做了海报和页面。", "围绕活动主题完成视觉风格、版式和物料设计，输出海报、横幅和页面素材，并根据反馈迭代版本。"),
        ("工程实践", "参与设备调试。", "配合完成设备参数记录、异常排查和测试结果整理，沉淀操作记录，减少后续重复排查成本。"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = value
    style_table(table, [1400, 3000, 4960])

    doc.add_heading("按岗位方向改写", level=1)
    add_bullets(doc, [
        "技术类岗位：突出技术方案、工具、难点、性能、质量和交付结果。",
        "运营类岗位：突出用户、内容、活动、数据、复盘和增长效果。",
        "产品类岗位：突出需求分析、用户调研、流程设计、原型、协作和上线反馈。",
        "市场销售类岗位：突出客户沟通、活动策划、转化结果、资源协调和目标达成。",
        "教育类岗位：突出备课、讲解、学生反馈、课程设计和教学结果。",
        "设计类岗位：突出作品集、审美逻辑、设计过程、用户反馈和落地场景。",
    ])

    doc.add_heading("可直接套用的句式", level=1)
    add_bullets(doc, [
        "围绕……目标，负责……工作，使用……方法，最终产出……。",
        "针对……问题，梳理……原因，提出……方案，并通过……验证效果。",
        "在……团队中承担……角色，协调……资源，保障……按期完成。",
        "基于……数据/反馈，发现……问题，调整……策略，使……得到改善。",
        "将……经验沉淀为……文档/模板/流程，便于后续复用。",
    ])

    doc.add_heading("智能体使用规则", level=1)
    add_bullets(doc, [
        "先追问用户真实经历，不直接代写不存在的成果。",
        "如果用户无法提供数据，可以使用“产出、反馈、效率、质量、复用价值”等非夸张结果。",
        "同一段经历可按不同岗位方向重写，但事实必须一致。",
        "面向学生的表达应自然可信，避免过度商业化或过度包装。",
    ])

    add_source_section(doc, ["国家大学生就业服务平台职位页", "中国公共招聘网", "全国就业公共服务平台"])
    return save_doc(doc, "实践经历表达示例.docx")


def build_interview_doc():
    doc = Document()
    set_doc_styles(doc)
    add_title(doc, "常见面试题与回答建议", "CyanCruise 简历面试知识库资料")

    doc.add_heading("适用场景", level=1)
    doc.add_paragraph(
        "本文用于支持 CyanCruise 智能体进行面试辅导、模拟问答和回答优化。"
        "内容覆盖大学生常见的自我介绍、经历追问、岗位动机、能力短板、压力问题和反问环节。"
    )

    doc.add_heading("面试前准备", level=1)
    add_numbers(doc, [
        "研究岗位要求：提取核心能力、职责、业务场景和硬性条件。",
        "准备个人证据：为每项核心能力匹配一个真实经历。",
        "梳理简历追问：对简历上每个项目准备背景、个人贡献、难点、结果和复盘。",
        "了解单位信息：业务、产品、行业、招聘流程、工作地点和培养机制。",
        "进行模拟练习：控制回答长度，避免背稿，练习追问和临场表达。",
    ])

    doc.add_heading("常见问题与回答建议", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "问题类型"
    table.rows[0].cells[1].text = "常见问题"
    table.rows[0].cells[2].text = "回答建议"
    rows = [
        ("自我介绍", "请做一个自我介绍。", "用 60 到 90 秒说明专业背景、核心经历、目标岗位匹配点和求职动机。"),
        ("岗位动机", "为什么选择这个岗位/行业？", "结合专业兴趣、经历证据和岗位理解，不要只说稳定、薪资或离家近。"),
        ("经历追问", "你在这个项目中具体做了什么？", "按经历讲述框架回答，重点讲个人任务、行动和结果。"),
        ("能力短板", "你觉得自己有什么不足？", "选择真实但可改进的短板，并说明已经采取的补强行动。"),
        ("压力问题", "如果任务很急、资源不足怎么办？", "说明优先级排序、沟通同步、风险预警和复盘改进。"),
        ("团队协作", "你和团队成员意见不一致怎么办？", "说明先对齐目标，再用事实和数据讨论方案，必要时推进小范围验证。"),
        ("职业规划", "未来三到五年怎么规划？", "表达学习成长、岗位深耕和阶段目标，避免空泛或频繁跳槽暗示。"),
        ("反问环节", "你有什么想问我们的？", "可问培养机制、团队协作方式、岗位前三个月重点任务、评价标准。"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = value
    style_table(table, [1500, 3000, 4860])

    doc.add_heading("回答质量检查", level=1)
    add_bullets(doc, [
        "是否正面回答了问题，而不是绕回简历背诵。",
        "是否包含具体经历，而不是只说性格和态度。",
        "是否能说明个人贡献，而不是只说团队成果。",
        "是否诚实，不夸大能力、不编造项目细节。",
        "是否有复盘意识，能说明学到了什么和下一次怎么改进。",
    ])

    doc.add_heading("不同面试形式提示", level=1)
    add_bullets(doc, [
        "单独面试：重视表达条理、简历一致性和岗位动机。",
        "群面：重视倾听、推进、总结、协作，不抢话也不沉默。",
        "技术/专业面试：准备基础概念、项目细节、问题排查和方案权衡。",
        "行为面试：用真实经历说明能力，不只给价值观口号。",
        "线上面试：提前测试网络、摄像头、麦克风、环境和材料。",
    ])

    doc.add_heading("智能体回答边界", level=1)
    add_bullets(doc, [
        "可以生成模拟面试问题、追问和回答改进建议。",
        "不能替用户编造经历、证书、成绩或项目结果。",
        "回答建议应保留用户真实语言风格，避免模板化过强。",
        "涉及具体单位、岗位流程和薪资时，应提示用户以招聘公告和面试通知为准。",
    ])

    add_source_section(doc, ["国家大学生就业服务平台", "国家大学生就业服务平台职位页", "教育部推广使用国家24365平台通知", "中国公共招聘网", "就业在线"])
    return save_doc(doc, "常见面试题与回答建议.docx")


def main():
    paths = [build_resume_guide(), build_experience_examples(), build_interview_doc()]
    zip_path = OUT_DIR / PACKAGE_NAME
    with ZipFile(zip_path, "w", ZIP_DEFLATED) as zf:
        for path in paths:
            zf.write(path, path.name)
    print("Generated:")
    for path in paths:
        print(path.resolve())
    print(zip_path.resolve())


if __name__ == "__main__":
    main()
