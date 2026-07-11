from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("docs/knowledge-base/cyancruise就业岗位知识库")
PACKAGE_NAME = "CyanCruise_就业岗位知识库_资料包.zip"


SOURCES = {
    "国家大学生就业服务平台": "https://www.ncss.cn/",
    "教育部推广使用国家24365平台通知": "https://www.moe.gov.cn/srcsite/A15/s3265/202204/t20220406_614117.html",
    "国家大学生就业服务平台职位页": "https://www.ncss.cn/student/m/jobs/index.html",
    "中国公共招聘网": "https://job.mohrss.gov.cn/",
    "全国就业公共服务平台": "https://www.12333.gov.cn/job/?channel=12333",
    "就业在线": "https://www.jobonline.cn/",
    "国家统计局2025年工资解读": "https://www.stats.gov.cn/sj/sjjd/202605/t20260515_1963706.html",
    "国家统计局2024年工资解读": "https://www.stats.gov.cn/xxgk/jd/sjjd2020/202505/t20250516_1959829.html",
    "职业分类大典情况说明": "https://www.scio.gov.cn/xwfb/bwxwfb/gbwfbh/rlzyhshbzb/202211/t20221111_618733.html",
    "2025年度新职业信息征集": "https://chrm.mohrss.gov.cn/%E4%BA%BA%E5%8A%9B%E8%B5%84%E6%BA%90%E7%A4%BE%E4%BC%9A%E4%BF%9D%E9%9A%9C%E9%83%A8%E5%BC%80%E5%B1%952025%E5%B9%B4%E5%BA%A6%E6%96%B0%E8%81%8C%E4%B8%9A%E4%BF%A1%E6%81%AF%E5%BE%81%E9%9B%86%E5%B7%A5/",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def style_table(table, widths=None):
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            if widths:
                set_cell_width(cell, widths[col_idx])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)
                for run in paragraph.runs:
                    run.font.name = "微软雅黑"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                    run.font.size = Pt(9.5)
            if row_idx == 0:
                set_cell_shading(cell, "F2F4F7")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def set_doc_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(20)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("1F4D78")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(subtitle)
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("555555")


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.1)
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.1)
        p.add_run(item)


def add_source_section(doc, source_names):
    doc.add_heading("资料来源", level=1)
    doc.add_paragraph("以下来源均为公开信息入口。岗位、薪资、招聘活动和行业变化具有时效性，应以平台实时信息和用人单位公告为准。")
    for name in source_names:
        doc.add_paragraph(f"{name}：{SOURCES[name]}", style="List Bullet")


def save_doc(doc, filename):
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUT_DIR / filename
    doc.save(path)
    return path


def build_disciplines_doc():
    doc = Document()
    set_doc_styles(doc)
    add_title(doc, "大学主要学科就业方向", "CyanCruise 就业岗位知识库资料")

    doc.add_heading("适用场景", level=1)
    doc.add_paragraph(
        "本文用于支持 CyanCruise 智能体根据学生专业背景，初步匹配就业方向、行业入口和岗位类型。"
        "内容面向大学主要学科门类，适合用于“我这个专业能做什么”“我适合哪些行业”的问答场景。"
    )

    doc.add_heading("学科门类与常见方向", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "学科门类"
    table.rows[0].cells[1].text = "常见就业方向"
    table.rows[0].cells[2].text = "需要重点评估"
    rows = [
        ("哲学类", "党政机关、研究机构、编辑出版、公共事务、文化传播、教育培训。", "表达能力、研究能力、写作能力、公共议题理解。"),
        ("经济学类", "银行、证券、保险、咨询、市场研究、数据分析、产业研究、政府与公共部门。", "数理基础、数据工具、财经知识、研究报告能力。"),
        ("法学类", "律所、企业法务、合规、知识产权、公共事务、公务员、基层治理。", "法律检索、文书写作、逻辑论证、资格考试路径。"),
        ("教育学类", "中小学教育、课程研发、教务运营、教育产品、心理与生涯辅导。", "教师资格、课堂表达、课程设计、学生沟通。"),
        ("文学类", "编辑、新媒体、品牌传播、文案策划、内容运营、外语服务、文化传媒。", "写作作品、语言能力、选题策划、平台运营。"),
        ("历史学类", "文博考古、档案管理、文化传播、教育、研究助理、地方文旅。", "资料整理、研究方法、讲解表达、专业继续深造。"),
        ("理学类", "科研助理、数据分析、实验技术、算法与建模、教师、技术支持。", "数学统计、实验技能、编程工具、读研必要性。"),
        ("工学类", "研发工程师、软件/硬件、制造工艺、质量管理、工程管理、技术销售。", "工程实践、项目经历、工具链、实习或竞赛。"),
        ("农学类", "农业技术推广、种业研发、食品检测、智慧农业、农产品运营、基层农技。", "实验/田间实践、政策理解、产业链视角。"),
        ("医学类", "临床、药企研发、医疗器械、公共卫生、医学运营、健康管理。", "执业资格、规培要求、学历门槛、伦理与合规。"),
        ("管理学类", "人力资源、运营、供应链、项目管理、市场营销、财务管理、产品经理。", "实习经历、数据分析、流程意识、沟通协调。"),
        ("艺术学类", "视觉设计、交互设计、影视制作、游戏美术、品牌设计、展陈策划。", "作品集、软件能力、审美表达、项目落地经验。"),
        ("交叉学科", "人工智能、智能制造、数字经济、低空经济、绿色低碳、智慧医疗等新方向。", "复合能力、行业理解、跨学科项目、学习速度。"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = value
    style_table(table, [1300, 4200, 3860])

    doc.add_heading("就业方向匹配方法", level=1)
    add_numbers(doc, [
        "先识别学生专业所属学科门类，再看专业课、项目、实习和兴趣是否指向某类行业。",
        "将方向分为“专业强相关”“能力可迁移”“需要补课转型”三类，不直接把学生限制在本专业。",
        "对学历门槛较高的方向，如科研、医学、法学、部分理工研发岗位，应提示升学或资格考试要求。",
        "对实践导向明显的方向，如运营、设计、工程、销售、供应链，应重点检查作品、项目和实习证据。",
        "根据用户目标城市、家庭约束和风险偏好，给出主路径与备选路径。",
    ])

    doc.add_heading("智能体追问模板", level=1)
    add_bullets(doc, [
        "你的专业、年级、成绩排名和最感兴趣的课程是什么？",
        "你更想走专业深耕、跨行业转型，还是先找稳定工作？",
        "你是否有项目、实习、竞赛、作品集、论文或社会实践经历？",
        "你是否接受读研、考证、考公、基层项目或去外地发展？",
        "你希望未来工作更偏研究、技术、管理、表达、服务还是创作？",
    ])

    doc.add_heading("回答边界", level=1)
    add_bullets(doc, [
        "可以给出学科到职业方向的初步匹配和能力补齐建议。",
        "不能保证某专业一定能进入某岗位，最终取决于岗位要求、个人能力和招聘周期。",
        "对医学、法律、教师等资格要求明显的方向，必须提醒用户核对执业资格或考试要求。",
    ])

    add_source_section(doc, ["国家大学生就业服务平台", "教育部推广使用国家24365平台通知", "职业分类大典情况说明", "2025年度新职业信息征集"])
    return save_doc(doc, "大学主要学科就业方向.docx")


def build_ability_doc():
    doc = Document()
    set_doc_styles(doc)
    add_title(doc, "不同岗位能力要求", "CyanCruise 就业岗位知识库资料")

    doc.add_heading("适用场景", level=1)
    doc.add_paragraph(
        "本文用于支持 CyanCruise 智能体识别用户目标岗位，并从通用能力、专业能力、作品或项目证据三个层面给出准备建议。"
        "岗位能力要求会随行业和企业变化，智能体应结合实时招聘信息二次核验。"
    )

    doc.add_heading("岗位类型与能力要求", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.rows[0].cells[0].text = "岗位类型"
    table.rows[0].cells[1].text = "代表岗位"
    table.rows[0].cells[2].text = "核心能力"
    table.rows[0].cells[3].text = "简历证据"
    rows = [
        ("技术研发", "软件开发、硬件、算法、实验研发、工程设计。", "专业基础、工具链、问题拆解、工程实现。", "项目、代码、实验数据、竞赛、论文、专利。"),
        ("产品与运营", "产品经理、内容运营、用户运营、活动运营、教育产品。", "用户理解、需求分析、数据意识、沟通推进。", "产品分析、运营活动、数据复盘、校园项目。"),
        ("市场与销售", "市场营销、品牌传播、销售管培、商务拓展。", "客户沟通、行业理解、方案表达、目标管理。", "实习业绩、活动策划、社团项目、调研报告。"),
        ("财经与咨询", "银行、证券、财务、审计、咨询、产业研究。", "财务基础、数据分析、研究写作、商业判断。", "证书、建模报告、实习、案例分析。"),
        ("法务与公共事务", "法务、合规、知识产权、公共事务、基层治理。", "法律检索、文书写作、逻辑论证、政策理解。", "案例分析、模拟法庭、实习、资格考试进度。"),
        ("教育与培训", "教师、课程研发、教务运营、学习规划。", "课程设计、表达讲解、学生沟通、教学反馈。", "试讲、教案、家教/支教、教师资格。"),
        ("医疗健康", "临床相关、药企、器械、公共卫生、医学运营。", "专业资格、医学知识、合规意识、服务沟通。", "规培/实习、科研、证书、病例或项目训练。"),
        ("设计与创作", "视觉设计、交互设计、影视、游戏美术、文案。", "审美判断、工具熟练、创意表达、交付能力。", "作品集、项目链接、比赛作品、商业稿件。"),
        ("制造与供应链", "工艺、质量、采购、计划、物流、设备管理。", "流程意识、现场问题解决、数据记录、安全规范。", "工厂实习、改善项目、工具证书、实践报告。"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = value
    style_table(table, [1400, 2200, 3000, 2760])

    doc.add_heading("通用能力模型", level=1)
    add_bullets(doc, [
        "学习能力：能快速理解新任务、新工具和新行业知识。",
        "表达能力：能清楚说明问题、方案、结果和复盘。",
        "协作能力：能在团队中分工、对齐、推进和反馈。",
        "数据意识：能用数据描述现状、衡量结果、支持判断。",
        "职业素养：守时、诚信、责任心、文档习惯和合规意识。",
    ])

    doc.add_heading("从岗位要求到准备计划", level=1)
    add_numbers(doc, [
        "收集 20 个目标岗位的招聘信息，提取高频技能、学历、经验和证书要求。",
        "把要求拆成“已经具备”“短期可补”“长期投入”“暂不匹配”四类。",
        "为每类核心能力匹配一条简历证据，例如项目、实习、作品、证书或课程成果。",
        "每周固定补强一个短板，并形成可展示成果，如报告、作品、代码、案例复盘。",
        "投递前用岗位要求逐项检查简历，不要只写课程和职责，要写任务、行动和结果。",
    ])

    doc.add_heading("智能体使用规则", level=1)
    add_bullets(doc, [
        "用户只说专业时，先推荐 3 到 5 个可能方向，并说明匹配理由。",
        "用户给出目标岗位时，先列能力差距，再给补强计划。",
        "用户缺少经历时，不要虚构经历，应建议用课程设计、竞赛、社会实践、作品集补证据。",
        "对岗位要求中的学历、证书、地域、实习时长，应提醒以招聘公告为准。",
    ])

    add_source_section(doc, ["国家大学生就业服务平台职位页", "中国公共招聘网", "全国就业公共服务平台", "就业在线", "职业分类大典情况说明"])
    return save_doc(doc, "不同岗位能力要求.docx")


def build_city_trend_doc():
    doc = Document()
    set_doc_styles(doc)
    add_title(doc, "城市与行业就业趋势", "CyanCruise 就业岗位知识库资料")

    doc.add_heading("适用场景", level=1)
    doc.add_paragraph(
        "本文用于支持 CyanCruise 智能体回答“去哪个城市发展”“哪个行业机会更多”“如何结合城市选择岗位”等问题。"
        "城市和行业趋势具有明显时效性，本文提供判断框架，具体岗位应通过官方就业平台和用人单位公告实时核验。"
    )

    doc.add_heading("官方趋势信号", level=1)
    add_bullets(doc, [
        "国家大学生就业服务平台和公共招聘平台可用于观察高校毕业生岗位供给、专场招聘和行业招聘活动。",
        "国家统计局工资数据可用于观察行业整体薪酬和岗位结构变化，但不能等同于单个城市或企业的实际 offer。",
        "人社部门新职业信息反映产业变化方向，例如先进制造、人工智能、低空经济、现代服务、康养托育等领域。",
        "地方公共就业平台和高校就业中心可补充区域性岗位、见习岗位、基层项目和校招信息。",
    ])

    doc.add_heading("城市选择维度", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "维度"
    table.rows[0].cells[1].text = "观察内容"
    table.rows[0].cells[2].text = "对学生的影响"
    rows = [
        ("产业匹配", "目标城市是否有对应行业集群、龙头企业、科研院所或公共岗位。", "决定岗位数量、成长空间和转岗机会。"),
        ("生活成本", "房租、通勤、消费、家庭支持、初期现金流压力。", "影响实际可支配收入和求职承压能力。"),
        ("学历门槛", "行业是否偏好研究生、资格证、实习或作品集。", "影响是否优先就业、考研或先实习。"),
        ("招聘节奏", "秋招、春招、地方招聘会、事业单位、基层项目时间。", "决定投递计划和备选路径。"),
        ("长期发展", "户籍、人才政策、家庭位置、行业稳定性、继续深造机会。", "影响三到五年的职业规划。"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = value
    style_table(table, [1600, 3900, 3860])

    doc.add_heading("行业趋势判断框架", level=1)
    add_numbers(doc, [
        "先看岗位需求：在国家大学生就业服务平台、中国公共招聘网等平台检索目标城市和目标行业。",
        "再看行业质量：关注岗位是否有清晰培养机制、合理薪酬、稳定用工和成长路径。",
        "再看个人匹配：把专业、技能、实习、证书、作品集与岗位要求逐项匹配。",
        "最后看风险：区分短期热门和长期适合，避免只因薪资高或城市热度高做选择。",
    ])

    doc.add_heading("重点关注方向", level=1)
    add_bullets(doc, [
        "数字经济相关：软件、数据、人工智能、网络安全、产品与运营等。",
        "先进制造相关：智能制造、电子信息、新能源汽车、工业软件、质量与工艺。",
        "绿色低碳相关：新能源、储能、环境治理、碳管理、电力系统。",
        "现代服务相关：金融、咨询、法律、教育、医疗健康、文化传媒。",
        "公共服务相关：基层项目、事业单位、教育医疗、公共管理和社区治理。",
    ])

    doc.add_heading("智能体回答边界", level=1)
    add_bullets(doc, [
        "可以基于公开趋势给出城市和行业选择框架。",
        "不能把全国平均工资直接当作用户能拿到的薪资。",
        "不能承诺某城市一定更好，应结合用户专业、家庭、风险偏好和岗位供给判断。",
        "如果涉及具体岗位薪资、招聘人数、截止时间，应提示用户查询实时招聘公告。",
    ])

    add_source_section(doc, ["国家大学生就业服务平台", "中国公共招聘网", "全国就业公共服务平台", "国家统计局2025年工资解读", "国家统计局2024年工资解读", "2025年度新职业信息征集"])
    return save_doc(doc, "城市与行业就业趋势.docx")


def main():
    paths = [build_disciplines_doc(), build_ability_doc(), build_city_trend_doc()]
    zip_path = OUT_DIR / PACKAGE_NAME
    with ZipFile(zip_path, "w", ZIP_DEFLATED) as zf:
        for path in paths:
            zf.write(path, path.name)
    print("Generated:")
    for path in paths:
        print(path.resolve())
    print(zip_path.resolve())


if __name__ == "__main__":
    main()
